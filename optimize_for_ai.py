import re
import sys
from pathlib import Path
from docx import Document
from markitdown import MarkItDown


def clean_docx(input_path: Path, temp_clean_path: Path):
    """
    Carga el archivo .docx y elimina encabezados y pies de página repetitivos
    que consumen tokens innecesarios en la IA.
    """
    doc = Document(input_path)

    # Eliminar encabezados y pies de página de todas las secciones
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        # Limpiar texto de los encabezados
        for p in section.header.paragraphs:
            p.text = ""
        # Limpiar texto de los pies de página
        for p in section.footer.paragraphs:
            p.text = ""

    doc.save(temp_clean_path)


def markdown_to_plain_text(md_text: str) -> str:
    """
    Convierte texto Markdown a texto plano real, eliminando toda la sintaxis
    de formato que no aporta significado semántico pero sí consume tokens
    (encabezados #, negritas **, tablas |, listas -, enlaces [texto](url), etc.).

    El objetivo es que el .txt resultante sea la representación más barata
    posible en tokens, conservando únicamente el contenido informativo.
    """
    text = md_text

    # Quitar bloques de código (```...```) dejando solo su contenido
    text = re.sub(r"```[a-zA-Z0-9]*\n?", "", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)

    # Imágenes: ![alt](url) -> alt (o se elimina si no hay alt)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)

    # Enlaces: [texto](url) -> texto (se descarta la URL, que no aporta significado)
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)

    # Encabezados: "### Título" -> "Título"
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)

    # Negrita / cursiva / tachado: **texto**, __texto__, *texto*, _texto_, ~~texto~~
    text = re.sub(r"(\*\*\*|___)(.+?)\1", r"\2", text)
    text = re.sub(r"(\*\*|__)(.+?)\1", r"\2", text)
    text = re.sub(r"(?<!\w)(\*|_)(.+?)\1(?!\w)", r"\2", text)
    text = re.sub(r"~~(.+?)~~", r"\1", text)

    # Líneas separadoras de tabla: |---|---|---|
    text = re.sub(r"^\s*\|?[\s:|-]+\|[\s:|-]*$", "", text, flags=re.MULTILINE)

    # Filas de tabla: quitar barras verticales y dejar los valores separados por espacio
    text = re.sub(r"^\s*\|", "", text, flags=re.MULTILINE)
    text = re.sub(r"\|\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\s*\|\s*", "  ", text)

    # Líneas horizontales: ---, ***, ___
    text = re.sub(r"^\s*([-*_])\1{2,}\s*$", "", text, flags=re.MULTILINE)

    # Blockquotes: "> texto" -> "texto"
    text = re.sub(r"^\s*>\s?", "", text, flags=re.MULTILINE)

    # Listas con viñetas: "- item" / "* item" / "+ item" -> "item"
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)

    # Listas numeradas: "1. item" -> "item"
    text = re.sub(r"^\s*\d+[.)]\s+", "", text, flags=re.MULTILINE)

    # Colapsar espacios múltiples dentro de una línea
    text = re.sub(r"[ \t]{2,}", " ", text)

    # Quitar espacios al final de cada línea
    text = re.sub(r"[ \t]+$", "", text, flags=re.MULTILINE)

    # Colapsar más de 2 saltos de línea seguidos en uno solo (ahorra tokens
    # y elimina espacios en blanco que no aportan significado)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def convert_to_ai_formats(input_file: str):
    file_path = Path(input_file).expanduser().resolve()

    # Si no existe en la carpeta actual, buscar en ~/Documents y ~/Downloads
    if not file_path.exists():
        search_dirs = [Path.home() / "Documents", Path.home() / "Downloads"]
        found = False
        for folder in search_dirs:
            candidate = folder / input_file
            if candidate.exists():
                file_path = candidate
                found = True
                break

        if not found:
            print(f"Error: No se encontró '{input_file}' en el directorio actual, ni en Documents ni en Downloads.")
            return

    print(f"Procesando: {file_path}")
    print("Limpiando el documento...")
    temp_clean_path = file_path.parent / f"temp_{file_path.name}"
    clean_docx(file_path, temp_clean_path)

    print("Convirtiendo a formatos para IA...")
    md = MarkItDown()
    result = md.convert(str(temp_clean_path))
    markdown_content = result.text_content

    # .md: se conserva la estructura completa (títulos, tablas, negritas...)
    # útil cuando la IA necesita entender jerarquía o relaciones en tablas.
    output_md = file_path.with_suffix('.md')
    with open(output_md, 'w', encoding='utf-8') as f:
        f.write(markdown_content)

    # .txt: texto plano real, sin sintaxis Markdown. Es el formato más barato
    # en tokens; ideal cuando solo importa el contenido, no el formato.
    plain_text = markdown_to_plain_text(markdown_content)
    output_txt = file_path.with_suffix('.txt')
    with open(output_txt, 'w', encoding='utf-8') as f:
        f.write(plain_text)

    if temp_clean_path.exists():
        temp_clean_path.unlink()

    # Reporte de tamaños para que el usuario vea el ahorro real de tokens
    md_chars = len(markdown_content)
    txt_chars = len(plain_text)
    saved_pct = (1 - txt_chars / md_chars) * 100 if md_chars else 0

    print(f"\n¡Listo! Archivos generados en {file_path.parent}:")
    print(f" - {output_md.name}  ({md_chars:,} caracteres, con formato Markdown)")
    print(f" - {output_txt.name}  ({txt_chars:,} caracteres, texto plano)")
    print(f"   Reducción de caracteres del .txt frente al .md: {saved_pct:.1f}%")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python optimize_for_ai.py mi_documento.docx")
    else:
        convert_to_ai_formats(sys.argv[1])